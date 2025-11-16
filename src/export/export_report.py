# src/export/export_report.py
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image
from reportlab.lib.styles import getSampleStyleSheet
import io, zipfile
import matplotlib.pyplot as plt

def build_report_pdf_docx(data, schedule_df):
    # build docx
    doc = Document()
    doc.add_heading("Báo cáo thẩm định", level=1)
    doc.add_paragraph("Thông tin định danh:")
    idf = data.get("identification", {})
    doc.add_paragraph(f"Họ tên: {idf.get('ten','')}")
    doc.add_paragraph(f"CCCD: {idf.get('cccd','')}")
    doc.add_paragraph("\nTóm tắt phương án:")
    fin = data.get("finance", {})
    doc.add_paragraph(f"Mục đích: {fin.get('muc_dich','')}")
    doc.add_paragraph(f"Số tiền vay: {fin.get('so_tien_vay','')}")
    # write docx to bytes
    docx_io = io.BytesIO()
    doc.save(docx_io)
    docx_io.seek(0)

    # build PDF (simple) with one chart
    pdf_io = io.BytesIO()
    docpdf = SimpleDocTemplate(pdf_io)
    styles = getSampleStyleSheet()
    elems = []
    elems.append(Paragraph("Báo cáo thẩm định", styles['Title']))
    elems.append(Paragraph(f"Họ tên: {idf.get('ten','')}", styles['Normal']))
    # chart
    if schedule_df is not None and not schedule_df.empty:
        fig, ax = plt.subplots()
        ax.plot(schedule_df["month"], schedule_df["payment"])
        ax.set_title("Nghĩa vụ trả nợ hàng tháng")
        img_io = io.BytesIO()
        fig.savefig(img_io, format="PNG", bbox_inches="tight")
        img_io.seek(0)
        elems.append(Image(img_io, width=400, height=200))
    docpdf.build(elems)
    pdf_io.seek(0)

    # zip both
    zip_io = io.BytesIO()
    with zipfile.ZipFile(zip_io, mode="w") as z:
        z.writestr("bao_cao.docx", docx_io.read())
        z.writestr("bao_cao.pdf", pdf_io.read())
    zip_io.seek(0)
    return zip_io.getvalue()
